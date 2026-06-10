"""
Konversi LAPORAN.md → LAPORAN.docx
Jalankan: python md_to_docx.py
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'BFBFBF')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run_with_inline(para, text: str, base_bold=False, base_italic=False):
    """Tambah run dengan parsing inline **bold**, *italic*, `code`."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = para.add_run(part[1:-1])
            r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
        else:
            if part:
                r = para.add_run(part)
                r.bold   = base_bold
                r.italic = base_italic


def set_page_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin   = Cm(4)
        section.right_margin  = Cm(3)


def apply_heading_style(para, level: int, doc: Document):
    colors = {
        1: '1F3864',   # biru tua
        2: '1F3864',
        3: '2E74B5',
        4: '2E74B5',
    }
    sizes  = {1: 18, 2: 15, 3: 13, 4: 12}
    color  = colors.get(level, '000000')
    size   = sizes.get(level, 11)

    for run in para.runs:
        run.font.color.rgb = RGBColor(
            int(color[0:2], 16),
            int(color[2:4], 16),
            int(color[4:6], 16),
        )
        run.font.size = Pt(size)
        run.bold      = True


# ── Parser utama ──────────────────────────────────────────────────────────────

def parse_and_build(md_path: str, docx_path: str):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    set_page_margins(doc)

    # Style default body
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    i = 0
    in_code_block   = False
    in_table        = False
    table_rows      = []
    ordered_counter = 0

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            in_table  = False
            table_rows = []
            return

        # Filter separator rows (---|---...)
        data_rows = [r for r in table_rows
                     if not re.match(r'^[\s|:\-]+$', r.replace('|', ''))]
        if not data_rows:
            in_table   = False
            table_rows = []
            return

        parsed = []
        for row in data_rows:
            cells = [c.strip() for c in row.strip().strip('|').split('|')]
            parsed.append(cells)

        max_cols = max(len(r) for r in parsed)
        tbl = doc.add_table(rows=len(parsed), cols=max_cols)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ri, row in enumerate(parsed):
            for ci in range(max_cols):
                cell  = tbl.cell(ri, ci)
                text  = row[ci] if ci < len(row) else ''
                # strip bold markers for cell text, handle inline
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                clean = re.sub(r'\*(.*?)\*',     r'\1', clean)
                clean = re.sub(r'`(.*?)`',       r'\1', clean)
                para  = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_cell_border(cell)
                if ri == 0:
                    set_cell_bg(cell, 'D6E4F0')
                    run = para.add_run(clean)
                    run.bold      = True
                    run.font.size = Pt(10)
                else:
                    run = para.add_run(clean)
                    run.font.size = Pt(10)
                    if ci == 0:
                        run.bold = True

        doc.add_paragraph()   # spasi setelah tabel
        in_table   = False
        table_rows = []

    while i < len(lines):
        raw  = lines[i]
        line = raw.rstrip('\n')

        # ── CODE BLOCK ────────────────────────────────────────────────────────
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                i += 1
                continue
            else:
                in_code_block = False
                i += 1
                continue

        if in_code_block:
            para = doc.add_paragraph()
            run  = para.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x20, 0x4A, 0x87)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(0)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'),   'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'),  'F2F2F2')
            para._p.get_or_add_pPr().append(shading)
            i += 1
            continue

        # ── TABLE ──────────────────────────────────────────────────────────────
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
            table_rows.append(line)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── HORIZONTAL RULE ───────────────────────────────────────────────────
        if re.match(r'^-{3,}$', line.strip()):
            doc.add_paragraph('─' * 60).paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ── HEADINGS ──────────────────────────────────────────────────────────
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level     = len(heading_match.group(1))
            text      = heading_match.group(2)
            # strip inline markup for heading text
            text_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text_clean = re.sub(r'\*(.*?)\*',     r'\1', text_clean)
            text_clean = re.sub(r'`(.*?)`',       r'\1', text_clean)
            style_name = f'Heading {min(level, 4)}'
            para = doc.add_heading(text_clean, level=min(level, 4))
            apply_heading_style(para, level, doc)
            i += 1
            continue

        # ── BLOCKQUOTE ────────────────────────────────────────────────────────
        if line.strip().startswith('>'):
            text = line.strip().lstrip('>').strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run(re.sub(r'\*\*(.*?)\*\*', r'\1',
                               re.sub(r'\*(.*?)\*', r'\1',
                               re.sub(r'`(.*?)`', r'\1', text))))
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.size = Pt(10)
            i += 1
            continue

        # ── ORDERED LIST ──────────────────────────────────────────────────────
        ol_match = re.match(r'^(\d+)\.\s+(.*)', line)
        if ol_match:
            ordered_counter += 1
            text = ol_match.group(2)
            para = doc.add_paragraph(style='List Number')
            para.paragraph_format.left_indent = Cm(1)
            add_run_with_inline(para, text)
            i += 1
            continue
        else:
            ordered_counter = 0

        # ── UNORDERED LIST ────────────────────────────────────────────────────
        ul_match = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if ul_match:
            indent = len(ul_match.group(1))
            text   = ul_match.group(2)
            para   = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Cm(1 + indent * 0.3)
            add_run_with_inline(para, text)
            i += 1
            continue

        # ── EMPTY LINE ────────────────────────────────────────────────────────
        if line.strip() == '':
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── NORMAL PARAGRAPH ─────────────────────────────────────────────────
        para = doc.add_paragraph()
        para.paragraph_format.space_after  = Pt(6)
        para.paragraph_format.first_line_indent = Cm(1)
        add_run_with_inline(para, line.strip())
        i += 1

    # flush tabel yang belum selesai
    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f'✓ Berhasil disimpan: {docx_path}')


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    import os
    base = os.path.dirname(os.path.abspath(__file__))
    parse_and_build(
        md_path   = os.path.join(base, 'LAPORAN.md'),
        docx_path = os.path.join(base, 'LAPORAN.docx'),
    )
